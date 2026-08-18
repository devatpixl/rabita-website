# -*- coding: utf-8 -*-
"""
Builds the Word copy of the SOP from the Markdown one.

The two drifted once already: the Markdown was updated and the Word file, which
is the copy anybody actually opens, was left a day behind. There is one source
now, and this makes the other from it.

    python docs/build-sop-docx.py
"""
import io
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "website-development-sop.md")
OUT = os.path.join(HERE, "Website Development SOP.docx")

BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x55, 0x55, 0x55)


def style(r, size=10.5, bold=False, color=BLACK, mono=False, caps=False, spacing=None):
    r.font.name = "Consolas" if mono else "Cambria"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    e = r._element
    e.rPr.rFonts.set(qn("w:eastAsia"), r.font.name)
    e.rPr.rFonts.set(qn("w:cs"), r.font.name)
    if caps:
        e.rPr.append(OxmlElement("w:caps"))
    if spacing:
        el = OxmlElement("w:spacing")
        el.set(qn("w:val"), str(spacing))
        e.rPr.append(el)


def para(doc, text="", size=10.5, bold=False, color=BLACK, before=0, after=6,
         line=1.25, indent=0, hanging=0, keep=False, mono=False, caps=False, spacing=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent:
        pf.left_indent = Cm(indent)
    if hanging:
        pf.first_line_indent = Cm(-hanging)
    pf.keep_with_next = keep
    if text:
        style(p.add_run(text), size, bold, color, mono, caps, spacing)
    return p


def inline(p, text, size=10.5, color=BLACK):
    """Bold for **x**, monospace for `x`, plain for the rest."""
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            style(p.add_run(part[2:-2]), size, True, color)
        elif part.startswith("`") and part.endswith("`"):
            style(p.add_run(part[1:-1]), size - 0.5, False, color, mono=True)
        else:
            style(p.add_run(part), size, False, color)


def table(doc, rows):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"
    for i, cells in enumerate(rows):
        row = t.add_row().cells
        for j, cell in enumerate(cells):
            row[j].text = ""
            p = row[j].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            inline(p, cell, size=9.5, color=BLACK if i == 0 else GREY)
            if i == 0:
                for r in p.runs:
                    r.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def build():
    md = io.open(SRC, encoding="utf-8").read().split("\n")

    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.4)

    i = 0
    while i < len(md):
        line = md[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # a table, collected whole
        if line.startswith("|"):
            rows = []
            while i < len(md) and md[i].startswith("|"):
                cells = [c.strip() for c in md[i].strip().strip("|").split("|")]
                if not all(set(c) <= set("- :") for c in cells):
                    rows.append(cells)
                i += 1
            table(doc, rows)
            continue

        # fenced code
        if line.startswith("```"):
            i += 1
            while i < len(md) and not md[i].startswith("```"):
                para(doc, md[i], size=9, color=GREY, after=0, line=1.15, indent=0.6, mono=True)
                i += 1
            i += 1
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        if line.startswith("### "):
            para(doc, line[4:], size=11.5, bold=True, before=12, after=4, keep=True)
        elif line.startswith("## "):
            para(doc, line[3:], size=13.5, bold=True, before=18, after=6, keep=True)
        elif line.startswith("# "):
            para(doc, line[2:], size=19, bold=True, before=0, after=2, line=1.05)
        elif line.startswith("- "):
            # a bullet, plus any wrapped continuation lines beneath it
            text = line[2:]
            while i + 1 < len(md) and md[i + 1].startswith("  ") and not md[i + 1].strip().startswith("-"):
                text += " " + md[i + 1].strip()
                i += 1
            p = para(doc, after=3, indent=0.7, hanging=0.35)
            style(p.add_run("\u2022\t"), 10.5, False, GREY)
            inline(p, text)
        else:
            p = para(doc, after=6)
            inline(p, line)
        i += 1

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
